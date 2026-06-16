import io
import uuid
from pathlib import Path
from typing import Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from ..config import get_settings

settings = get_settings()


# Prefix every key under `psydict/` so we never collide with other apps that
# share the same bucket (e.g. bsl-fotos hosts kbnet/, fotos/, firmas/, …).
SPACES_PREFIX = "psydict"


def _upload_to_spaces(data: bytes, filename: str, content_type: str) -> str:
    """Upload file to DigitalOcean Spaces and return public URL."""
    key = f"{SPACES_PREFIX}/{filename}"
    if not settings.SPACES_KEY:
        local_path = Path("/tmp") / key
        local_path.parent.mkdir(parents=True, exist_ok=True)
        local_path.write_bytes(data)
        return f"{settings.BASE_URL}/static/{key}"
    import boto3
    from botocore.client import Config
    s3 = boto3.client(
        "s3",
        region_name=settings.SPACES_REGION,
        endpoint_url=settings.SPACES_ENDPOINT,
        aws_access_key_id=settings.SPACES_KEY,
        aws_secret_access_key=settings.SPACES_SECRET,
        config=Config(signature_version="s3v4"),
    )
    s3.put_object(
        Bucket=settings.SPACES_BUCKET,
        Key=key,
        Body=data,
        ContentType=content_type,
        ACL="public-read",
    )
    # Virtual-host style: https://<bucket>.<region>.digitaloceanspaces.com/<key>
    return f"https://{settings.SPACES_BUCKET}.{settings.SPACES_REGION}.digitaloceanspaces.com/{key}"


def _escape_html(s: str) -> str:
    """ReportLab Paragraph parses limited HTML; escape so '&', '<', '>' inside
    Claude's prose don't break rendering."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_pdf(content: dict[str, Any], project_id: str) -> str:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("APATitle", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=12, fontName="Times-Roman")
    heading1 = ParagraphStyle("APAHeading1", parent=styles["Heading1"], fontSize=12, fontName="Times-Bold", alignment=TA_CENTER)
    body = ParagraphStyle("APABody", parent=styles["Normal"], fontSize=12, fontName="Times-Roman", leading=24)

    story = []
    story.append(Paragraph(content.get("title", "Untitled"), title_style))
    story.append(Spacer(1, 0.5 * inch))

    sections_rendered = 0
    for section_key in ["abstract", "introduction", "method", "results", "discussion", "references"]:
        if section_key in content and content[section_key]:
            story.append(Paragraph(section_key.capitalize(), heading1))
            story.append(Spacer(1, 0.25 * inch))
            story.append(Paragraph(_escape_html(content[section_key]), body))
            story.append(Spacer(1, 0.25 * inch))
            sections_rendered += 1

    # Resilience: if Claude's structured output failed (parse error,
    # truncation, missing keys), surface the raw text so the user gets
    # SOMETHING usable instead of a near-empty PDF with just a title.
    if sections_rendered == 0 and content.get("raw"):
        story.append(Paragraph("Generated content (unstructured)", heading1))
        story.append(Spacer(1, 0.25 * inch))
        # Reportlab Paragraph handles \n as newline only inside <br/>, so split.
        for chunk in str(content["raw"]).split("\n\n"):
            if chunk.strip():
                story.append(Paragraph(_escape_html(chunk), body))
                story.append(Spacer(1, 0.1 * inch))

    doc.build(story)
    data = buffer.getvalue()
    filename = f"documents/{project_id}_{uuid.uuid4().hex[:8]}.pdf"
    return _upload_to_spaces(data, filename, "application/pdf")


def generate_docx(content: dict[str, Any], project_id: str) -> str:
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading(content.get("title", "Untitled"), level=0)

    sections_rendered = 0
    for section_key in ["abstract", "introduction", "method", "results", "discussion", "references"]:
        if section_key in content and content[section_key]:
            doc.add_heading(section_key.capitalize(), level=1)
            doc.add_paragraph(content[section_key])
            sections_rendered += 1

    if sections_rendered == 0 and content.get("raw"):
        doc.add_heading("Generated content (unstructured)", level=1)
        for chunk in str(content["raw"]).split("\n\n"):
            if chunk.strip():
                doc.add_paragraph(chunk)

    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    filename = f"documents/{project_id}_{uuid.uuid4().hex[:8]}.docx"
    return _upload_to_spaces(data, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
